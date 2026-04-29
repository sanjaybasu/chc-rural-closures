"""
Build a properly styled JHCPU DOCX from manuscript.md using python-docx.
- Cambria Math font, black, 1.5 line spacing throughout
- Real superscripts for citation markers (<sup>...</sup>)
- Hard line breaks between Objective / Methods / Results / Conclusions
- Embedded figures (not external links)
- Embedded tables rendered from CSVs

Outputs:
  manuscript/manuscript_jhcpu.docx
  manuscript/title_page.docx
  manuscript/cover_letter.docx
"""
from pathlib import Path
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/sanjaybasu/waymark-local/notebooks/chc-rural-closures")
MS = ROOT/"manuscript"
RES = ROOT/"results"
FIG = ROOT/"figures"

FONT_NAME = "Cambria Math"
BLACK = RGBColor(0, 0, 0)

def set_run_style(run, *, size=11, bold=False, italic=False, superscript=False):
    run.font.name = FONT_NAME
    run.font.color.rgb = BLACK
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if superscript:
        run.font.superscript = True
    # ensure complex script font is also Cambria Math
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)

def set_para_spacing(p, *, line_spacing=1.5, space_after=6):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)

def add_inline_runs(p, text, base_size=11, bold=False, italic=False):
    """Parse markdown-ish bold (**x**), italic (*x*), and <sup>...</sup>; emit runs."""
    # split keeping delimiters; handle in order: <sup>, **, *
    pattern = re.compile(r'(<sup>.*?</sup>|\*\*.*?\*\*|\*[^*]+\*)')
    parts = pattern.split(text)
    for part in parts:
        if not part: continue
        if part.startswith("<sup>") and part.endswith("</sup>"):
            inner = part[5:-6]
            r = p.add_run(inner)
            set_run_style(r, size=base_size, bold=bold, italic=italic, superscript=True)
        elif part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            r = p.add_run(inner)
            set_run_style(r, size=base_size, bold=True, italic=italic)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            inner = part[1:-1]
            r = p.add_run(inner)
            set_run_style(r, size=base_size, bold=bold, italic=True)
        else:
            r = p.add_run(part)
            set_run_style(r, size=base_size, bold=bold, italic=italic)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_para_spacing(p, line_spacing=1.5, space_after=8)
    sizes = {1: 14, 2: 12, 3: 11}
    r = p.add_run(text)
    set_run_style(r, size=sizes.get(level, 11), bold=True)
    return p

def add_para(doc, text, *, base_size=11, bold=False):
    p = doc.add_paragraph()
    set_para_spacing(p)
    add_inline_runs(p, text, base_size=base_size, bold=bold)
    return p

def add_image(doc, path, width_inches=6.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p)
    p.add_run().add_picture(str(path), width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        set_para_spacing(cap)
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline_runs(cap, caption, base_size=10)

def add_table_from_df(doc, df, *, col_widths=None):
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        set_para_spacing(p, line_spacing=1.15, space_after=2)
        r = p.add_run(str(c))
        set_run_style(r, size=10, bold=True)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, val in enumerate(row.values):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            set_para_spacing(p, line_spacing=1.15, space_after=2)
            r = p.add_run(str(val))
            set_run_style(r, size=10)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t

# -------------------------------------------------------------------
# MAIN MANUSCRIPT
# -------------------------------------------------------------------
def build_manuscript():
    doc = Document()
    # base style: Cambria Math, 11 pt, 1.5 spacing
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    md = (MS/"manuscript.md").read_text()

    # --- Header block ---
    title = re.search(r"^# (.+)", md, flags=re.MULTILINE).group(1)
    p = doc.add_paragraph(); set_para_spacing(p, space_after=12)
    r = p.add_run(title); set_run_style(r, size=14, bold=True)

    add_para(doc, "**Article type:** Brief Communication")
    add_para(doc, "**Running head:** CHCs and rural hospital closures")
    add_para(doc, "**Author:** Sanjay Basu, MD, PhD")
    add_para(doc, "**Affiliations:** Department of Medicine, University of California, San Francisco; Waymark, San Francisco, CA")
    add_para(doc, "**Correspondence:** sanjay.basu@ucsf.edu")

    # --- Abstract ---
    add_heading(doc, "Abstract", level=2)
    abs_block = md.split("## Abstract")[1].split("---")[0]
    abs_lines = [l.strip() for l in abs_block.strip().splitlines() if l.strip()]
    for line in abs_lines:
        if line.startswith("**Key words"):
            add_para(doc, line)
        else:
            add_para(doc, line)

    # --- Body sections ---
    body = md.split("## (Introduction)")[1].split("## References")[0]
    # Walk sections
    sections = re.split(r"^## (.+)$", body, flags=re.MULTILINE)
    # sections[0] is the introduction text (no heading), then alternating heading,text...
    intro_text = sections[0].strip()
    for para in re.split(r"\n\n+", intro_text):
        if para.strip(): add_para(doc, para.strip())
    for i in range(1, len(sections), 2):
        heading = sections[i].strip()
        text = sections[i+1].strip() if i+1 < len(sections) else ""
        add_heading(doc, heading, level=2)
        for para in re.split(r"\n\n+", text):
            para = para.strip()
            if not para: continue
            # detect bold inline subheadings like "**Patient volume.**"
            add_para(doc, para)

    # --- Acknowledgments ---
    if "## Acknowledgments" in md:
        ack = md.split("## Acknowledgments")[1].split("##")[0].strip()
        add_heading(doc, "Acknowledgments", level=2)
        for para in re.split(r"\n\n+", ack):
            if para.strip(): add_para(doc, para.strip())

    # --- References ---
    add_heading(doc, "References", level=2)
    refs = md.split("## References")[1].split("---")[0].strip()
    for line in refs.splitlines():
        line = line.strip()
        if not line: continue
        if re.match(r"^\d+\.\s", line):
            add_para(doc, line)

    # --- Figures & Tables ---
    doc.add_page_break()
    add_heading(doc, "Figures and Tables", level=2)

    add_para(doc, "**Figure 1.** Forest plot of Callaway–Sant'Anna simple ATTs across all 11 pre-specified outcomes, exposed vs. never-exposed CHCs, 2014–2024. Blue intervals indicate statistical significance at 5% (raw p<0.05); intervals are 95% confidence intervals.")
    add_image(doc, FIG/"fig2_forest_all_outcomes.png", width_inches=6.5)

    add_para(doc, "**Figure 2.** Event-study Callaway–Sant'Anna ATT estimates by relative period for all 11 outcomes. Red filled circles indicate event-time-specific 95% pointwise confidence bands excluding zero. Reference period: year before exposure.")
    add_image(doc, FIG/"fig3_event_panels.png", width_inches=6.5)

    # Table 1
    add_para(doc, "**Table 1.** Baseline (2014) characteristics of CHCs by closure exposure.")
    t1 = pd.read_csv(RES/"table1_baseline.csv")
    add_table_from_df(doc, t1)

    # Table 2: simple ATTs labeled
    LABEL_ORDER = [
        ("log_n_total", "log(Total patients)"),
        ("share_uninsured", "Uninsured share (pp)"),
        ("share_medicaid", "Medicaid share (pp)"),
        ("htn_control_pct", "Hypertension control (pp)"),
        ("dm_poor_control_pct", "Diabetes A1c>9% (pp)"),
        ("imm_child", "Childhood immunization (pp)"),
        ("pap_screen", "Cervical cancer screening (pp)"),
        ("crc_screen", "Colorectal cancer screening (pp)"),
        ("bmi_adult", "Adult BMI follow-up (pp)"),
        ("tobacco", "Tobacco assess+intervention (pp)"),
        ("depr_screen", "Depression screen+follow-up (pp)"),
    ]
    src = pd.read_csv(RES/"cs2_simple_all.csv").set_index("outcome")
    rows = []
    for k, lab in LABEL_ORDER:
        r = src.loc[k]
        scale = 1.0 if k=="log_n_total" else 100.0
        att = float(r["att"])*scale; lo = float(r["ci_lo"])*scale; hi = float(r["ci_hi"])*scale
        rows.append([lab, f"{att:+.3f}", f"{lo:+.3f} to {hi:+.3f}"])
    t2 = pd.DataFrame(rows, columns=["Outcome", "Simple ATT", "95% CI"])
    add_para(doc, "**Table 2.** Callaway–Sant'Anna simple ATT estimates, exposed vs. never-exposed CHCs (n=1,642; 386 exposed). For log(Total patients), values are log units; for all other outcomes, percentage points (pp).")
    add_table_from_df(doc, t2)

    # Table 3: multiple-testing
    mt = pd.read_csv(RES/"multiple_testing_adjusted.csv")
    order_map = {k:i for i,(k,_) in enumerate(LABEL_ORDER)}
    mt["sort"] = mt["outcome"].map(order_map)
    mt = mt.sort_values("sort")
    t3 = mt[["label","p_raw","p_bonf","p_holm","p_BH","p_BY","p_RW"]].copy()
    t3.columns = ["Outcome", "Raw p", "Bonferroni", "Holm", "BH (FDR)", "BY (FDR)", "Romano-Wolf"]
    for c in t3.columns[1:]:
        t3[c] = t3[c].apply(lambda x: f"{float(x):.4f}" if float(x) >= 0.0001 else "<.0001")
    add_para(doc, "**Table 3.** Multiple-testing-adjusted p-values across 11 pre-specified outcomes.")
    add_table_from_df(doc, t3)

    out = MS/"manuscript_jhcpu.docx"
    doc.save(out)
    print(f"saved {out}  ({out.stat().st_size:,} bytes)")

# -------------------------------------------------------------------
# TITLE PAGE
# -------------------------------------------------------------------
def build_simple(md_path, out_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    md = md_path.read_text()
    for line in md.splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(); set_para_spacing(p, space_after=12)
            r = p.add_run(line[2:]); set_run_style(r, size=14, bold=True)
        elif line.startswith("## "):
            p = doc.add_paragraph(); set_para_spacing(p, space_after=8)
            r = p.add_run(line[3:]); set_run_style(r, size=12, bold=True)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_para_spacing(p)
            add_inline_runs(p, line[2:])
        else:
            add_para(doc, line)
    doc.save(out_path)
    print(f"saved {out_path}  ({out_path.stat().st_size:,} bytes)")

if __name__ == "__main__":
    build_manuscript()
    build_simple(MS/"title_page.md",      MS/"title_page.docx")
    build_simple(MS/"cover_letter.md",    MS/"cover_letter.docx")
    build_simple(MS/"author_contacts.md", MS/"author_contacts.docx")
